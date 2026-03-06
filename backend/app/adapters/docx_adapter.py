"""
Layer A — DOCX Adapter (python-docx)

Extracts paragraphs, runs, styles, numbering, tables, and footnotes.
Paragraph style names from Word are highly reliable structural signals,
so they are preserved verbatim in source_provenance for the heuristic layer.
"""

from __future__ import annotations

import hashlib
import logging
from pathlib import Path
from typing import Any, Dict, List, Optional

from ..pipeline.canonical_model import Block, CanonicalDoc, SourceType, Span

logger = logging.getLogger(__name__)


class DOCXAdapter:
    """Ingest a DOCX file and produce a CanonicalDoc."""

    def ingest(self, file_path: Path) -> CanonicalDoc:
        try:
            import docx  # type: ignore
            from docx.oxml.ns import qn  # type: ignore
        except ImportError as e:
            raise RuntimeError("python-docx not installed. Run: pip install python-docx") from e

        doc_id = _file_hash(file_path)
        word_doc = docx.Document(str(file_path))

        metadata: Dict[str, Any] = {}
        core = word_doc.core_properties
        if core:
            for attr in ("title", "subject", "author", "description", "keywords", "language"):
                val = getattr(core, attr, None)
                if val:
                    metadata[attr] = str(val)

        blocks: List[Block] = []

        # Extract main body paragraphs
        for para in word_doc.paragraphs:
            block = _para_to_block(para)
            if block:
                blocks.append(block)

        # Extract tables
        for table in word_doc.tables:
            table_blocks = _table_to_blocks(table)
            blocks.extend(table_blocks)

        # Extract footnotes if available
        footnote_blocks = _extract_footnotes(word_doc)
        blocks.extend(footnote_blocks)

        return CanonicalDoc(
            doc_id=doc_id,
            source_type=SourceType.DOCX,
            metadata=metadata,
            blocks=blocks,
        )


# ---------------------------------------------------------------------------
# Paragraph extraction
# ---------------------------------------------------------------------------

def _para_to_block(para: Any) -> Optional[Block]:
    """Convert a python-docx Paragraph to a Block."""
    try:
        from docx.oxml.ns import qn  # type: ignore
    except ImportError:
        return None

    full_text = para.text
    if not full_text.strip():
        return None

    style_name = para.style.name if para.style else ""

    # Numbering info
    num_props = _get_numbering_props(para)

    spans: List[Span] = []
    for run in para.runs:
        text = run.text
        if not text:
            continue
        font = run.font
        font_size = float(font.size.pt) if font and font.size else None
        color_rgb = None
        if font and font.color and font.color.type is not None:
            try:
                rgb = font.color.rgb
                if rgb:
                    color_rgb = f"#{str(rgb)}"
            except Exception:
                pass

        spans.append(
            Span(
                text=text,
                bold=bool(run.bold),
                italic=bool(run.italic),
                font_name=font.name if font else None,
                font_size=font_size,
                color=color_rgb,
            )
        )

    if not spans:
        # para.text exists but no runs — unlikely; create a plain span
        spans = [Span(text=full_text)]

    prov: Dict[str, Any] = {
        "docx_style": style_name,
    }
    if num_props:
        prov.update(num_props)

    return Block(
        text=full_text,
        spans=spans,
        source_provenance=prov,
    )


def _get_numbering_props(para: Any) -> Optional[Dict[str, Any]]:
    """
    Extract numbering (list) properties from a paragraph's XML.
    Returns dict with numbering_id, numbering_depth, numbering_format or None.
    """
    try:
        from docx.oxml.ns import qn  # type: ignore
    except ImportError:
        return None

    pPr = para._p.find(qn("w:pPr"))
    if pPr is None:
        return None

    numPr = pPr.find(qn("w:numPr"))
    if numPr is None:
        return None

    ilvl = numPr.find(qn("w:ilvl"))
    numId = numPr.find(qn("w:numId"))

    depth = int(ilvl.get(qn("w:val"), 0)) if ilvl is not None else 0
    num_id = numId.get(qn("w:val"), "0") if numId is not None else "0"

    # Try to get numFmt from the numbering part
    num_format = _resolve_num_format(para, num_id, depth)

    return {
        "numbering_id": num_id,
        "numbering_depth": depth,
        "numbering_format": num_format,
    }


def _resolve_num_format(para: Any, num_id: str, depth: int) -> str:
    """
    Look up the numFmt for this numId+ilvl in the document's numbering.xml.
    Returns 'bullet', 'decimal', 'lowerLetter', etc. or 'unknown'.
    """
    try:
        from docx.oxml.ns import qn  # type: ignore
        doc = para._p.getroottree().getroot()
        # Walk up to document root to find numbering part reference
        # This is a simplified lookup — full resolution requires traversing abstractNumId
        part = para.part
        if not hasattr(part, "numbering_part") or part.numbering_part is None:
            return "bullet"
        num_part = part.numbering_part
        for num_elem in num_part._element.findall(qn("w:num")):
            if num_elem.get(qn("w:numId")) == num_id:
                abs_id_elem = num_elem.find(qn("w:abstractNumId"))
                if abs_id_elem is None:
                    break
                abs_id = abs_id_elem.get(qn("w:val"))
                for abs_num in num_part._element.findall(qn("w:abstractNum")):
                    if abs_num.get(qn("w:abstractNumId")) == abs_id:
                        for lvl in abs_num.findall(qn("w:lvl")):
                            if lvl.get(qn("w:ilvl")) == str(depth):
                                fmt_elem = lvl.find(qn("w:numFmt"))
                                if fmt_elem is not None:
                                    return fmt_elem.get(qn("w:val"), "bullet")
    except Exception:
        pass
    return "bullet"


# ---------------------------------------------------------------------------
# Table extraction
# ---------------------------------------------------------------------------

def _table_to_blocks(table: Any) -> List[Block]:
    """Convert a python-docx Table to a list of table-row Blocks."""
    blocks: List[Block] = []
    for row_idx, row in enumerate(table.rows):
        cell_texts = []
        cell_spans: List[Span] = []
        for cell in row.cells:
            cell_text = cell.text.strip()
            cell_texts.append(cell_text)
            # Check if any paragraph in the cell is bold (header signal)
            is_bold = any(
                any(run.bold for run in para.runs)
                for para in cell.paragraphs
                if para.runs
            )
            cell_spans.append(Span(text=cell_text, bold=is_bold))

        row_text = " | ".join(cell_texts)
        if not row_text.strip():
            continue

        # Heuristic: first row or all-bold cells → mark as header candidate
        all_bold = all(s.bold for s in cell_spans if s.text)
        is_header = row_idx == 0 and all_bold

        blocks.append(
            Block(
                text=row_text,
                spans=cell_spans,
                source_provenance={
                    "block_type": "table_row",
                    "row_index": row_idx,
                    "is_header_row": is_header,
                    "cell_count": len(cell_texts),
                },
            )
        )

    return blocks


# ---------------------------------------------------------------------------
# Footnote extraction
# ---------------------------------------------------------------------------

def _extract_footnotes(word_doc: Any) -> List[Block]:
    """
    Attempt to extract footnotes from the document's footnotes part.
    python-docx does not expose footnotes as a high-level API,
    so we access the XML directly.
    Returns an empty list if footnotes are unavailable.
    """
    blocks: List[Block] = []
    try:
        from docx.oxml.ns import qn  # type: ignore
        part = word_doc.part
        if not hasattr(part, "footnotes_part") or part.footnotes_part is None:
            return blocks
        fn_part = part.footnotes_part
        for fn in fn_part._element.findall(qn("w:footnote")):
            fn_id = fn.get(qn("w:id"), "")
            # Skip separator footnotes (id -1, 0)
            if fn_id in ("-1", "0"):
                continue
            texts = []
            for t_elem in fn.iter(qn("w:t")):
                t = (t_elem.text or "").strip()
                if t:
                    texts.append(t)
            fn_text = " ".join(texts)
            if fn_text:
                blocks.append(
                    Block(
                        text=fn_text,
                        spans=[Span(text=fn_text)],
                        source_provenance={
                            "block_type": "footnote",
                            "footnote_id": fn_id,
                        },
                    )
                )
    except Exception as exc:
        logger.debug("Footnote extraction failed (non-fatal): %s", exc)

    return blocks


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _file_hash(path: Path) -> str:
    h = hashlib.sha1()
    with open(path, "rb") as f:
        for chunk in iter(lambda: f.read(65536), b""):
            h.update(chunk)
    return h.hexdigest()[:16]
